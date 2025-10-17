from tracemalloc import start
from django.http import JsonResponse
from datetime import datetime, timedelta
from weasyprint import HTML
from django.template.loader import get_template
from django.http import HttpResponse
from django.db.models import Q, F, Aggregate, CharField
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement

from project_management.models import DailyProjectUpdate, Project, ProjectHour
from project_management.serializers import ProjectDetailSerializer, ProjectListSerializer
from project_management.utils.auto_client_weekly_report import ClientWeeklyUpdate
from rest_framework.generics import RetrieveAPIView, ListAPIView
from drf_yasg.utils import swagger_auto_schema
from drf_yasg import openapi

from django_filters import rest_framework as django_filters
from rest_framework import filters
from rest_framework.pagination import PageNumberPagination
from django_filters.rest_framework import DjangoFilterBackend



class GroupConcat(Aggregate):
    function = "GROUP_CONCAT"
    template = "%(function)s(%(distinct)s%(expressions)s%(ordering)s%(separator)s)"

    def __init__(
        self, expression, distinct=False, ordering=None, separator=",", **extra
    ):
        super().__init__(
            expression,
            distinct="DISTINCT " if distinct else "",
            ordering=" ORDER BY %s" % ordering if ordering is not None else "",
            separator=' SEPARATOR "%s"' % separator,
            output_field=CharField(),
            **extra,
        )


# Create your views here.
def get_this_week_hour(request, project_id, hour_date):
    project_hour_id = request.GET.get("project_hour_id")
    manager_id = request.user.employee.id
    if project_hour_id:
        manager_id = ProjectHour.objects.get(id=project_hour_id).manager_id

    # employee = (
    #     Employee.objects.filter(active=True, project_eligibility=True)
    #     .annotate(
    #         total_hour=Coalesce(
    #             Sum(
    #                 "dailyprojectupdate_employee__hours",
    #                 filter=Q(
    #                     dailyprojectupdate_employee__project=project_id,
    #                     dailyprojectupdate_employee__manager=manager_id,
    #                     dailyprojectupdate_employee__status="approved",
    #                     dailyprojectupdate_employee__created_at__date__lte=hour_date,
    #                     dailyprojectupdate_employee__created_at__date__gte=hour_date
    #                     - timedelta(days=6),
    #                 ),
    #             ),
    #             0.0,
    #         ),
    #         update=F("dailyprojectupdate_employee__updates_json"),
    #         update_id=GroupConcat(F("dailyprojectupdate_employee__id")),
    #     )
    #     .exclude(total_hour=0.0)
    #     .values("id", "full_name", "total_hour", "update", "update_id")
    # )
    q_obj = Q(
        project=project_id,
        manager=manager_id,
        status="approved",
        created_at__date__lte=hour_date,
        created_at__date__gte=hour_date - timedelta(days=6),
    )
    d = DailyProjectUpdate.objects.filter(
        q_obj,
        employee__active=True,
        employee__project_eligibility=True,
    )
    employee = (
        DailyProjectUpdate.objects.filter(
            q_obj,
            employee__active=True,
            employee__project_eligibility=True,
        )
        .annotate(
            full_name=F("employee__full_name"),
        )
        .exclude(hours=0.0)
        .values(
            "id",
            "created_at",
            "employee_id",
            "full_name",
            "hours",
            "update",
            "updates_json",
        )
    )

    totalHours = sum(hour["hours"] for hour in employee)

    employeeList = filter(lambda emp: emp["employee_id"], employee)

    data = {
        "manager_id": manager_id,
        "weekly_hour": list(employeeList),
        "total_project_hours": totalHours,
    }

    return JsonResponse(data)


def slack_callback(request):
    code = request.GET.get("code")
    state = request.GET.get("state")

    return JsonResponse({"code": code, "state": state})


def generate_weekly_update_word(response, data: dict):
    # Create a new Word document
    project = data.get("project")
    start_date = data.get("start_date")
    end_date = data.get("end_date")
    document = Document()

    # Add title
    title = document.add_paragraph()
    title_run = title.add_run(f"Weekly Development Update: {project.title}")
    title_run.bold = True
    title_run.font.size = Pt(14)
    # title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add date
    date_paragraph = document.add_paragraph(f"Date: {end_date} – {start_date}")
    # date_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add a line break
    # document.add_paragraph("")

    # Iterate through tasks and subtasks
    tasks = data.get("reports")
    for item in tasks.get("tasks"):
        # Add section title
        section_title = document.add_paragraph(f"• {item['feature']}")
        # section_title_run = section_title.add_run(f"• {item['feature']}")
        # section_title_run.bold = True
        # section_title_run.font.size = Pt(12)

        # Add subtasks as bullet points
        for task in item["subtasks"]:
            bullet_point = document.add_paragraph()
            bullet_point.add_run(f"• {task}").font.size = Pt(
                11
            )  # Add bullet point and task text

        # Add a line break after each section
        # document.add_paragraph("")

    # Save the document
    document.save(response)


def generate_client_weekly_report(request, project_id, hour_date):
    project_hour_id = request.GET.get("project_hour_id")
    doc_type = request.GET.get("doc_type")
    manager_id = request.user.employee.id
    if project_hour_id:
        manager_id = ProjectHour.objects.get(id=project_hour_id).manager_id

    q_obj = Q(
        project_id=project_id,
        status="approved",
        created_at__date__lte=hour_date,
        created_at__date__gte=hour_date - timedelta(days=6),
    )
    if not request.user.is_superuser:
        q_obj = q_obj | Q(manager=manager_id)
    employee = (
        DailyProjectUpdate.objects.filter(
            q_obj,
            employee__active=True,
            employee__project_eligibility=True,
        )
        .annotate(
            full_name=F("employee__full_name"),
        )
        .exclude(hours=0.0)
        .values(
            "id",
            "update",
            "updates_json",
        )
    )
    update = "\n\n".join([i["updates_json"][0][0] for i in employee])
    if not update:
        return JsonResponse({"status": 404, "state": "No Update Found"})
    open_ai_res = ClientWeeklyUpdate(update)
    data = open_ai_res.chat()
    print(data)
    template_name = "admin/client_weekly_report.html"
    template = get_template(template_name)
    context = {
        "reports": data,
        "project": Project.objects.get(id=project_id),
        "start_date": hour_date,
        "end_date": hour_date - timedelta(days=6),
    }
    html_content = template.render(context)
    if doc_type == "docx".lower():
        response = HttpResponse(
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        response["Content-Disposition"] = 'attachment; filename="weekly_update.docx"'
        generate_weekly_update_word(response, context)
        return response

    else:

        # Generate PDF
        html = HTML(string=html_content)
        pdf_file = html.write_pdf()
        return HttpResponse(pdf_file, content_type="application/pdf")



def generate_pdf(request, *args, **kwargs):
    data = request.POST
    open_ai_res = ClientWeeklyUpdate(data["update"])
    response = open_ai_res.chat()
    print(response)
    hour_date_str = request.POST.get("hour_date")

    # Convert the string to a datetime object
    # Assuming the date format is 'YYYY-MM-DD'
    hour_date = datetime.strptime(hour_date_str, "%Y-%m-%d")
    start_date = hour_date.strftime("%d %B, %Y")
    end_data = hour_date - timedelta(days=5)
    template_name = "admin/client_weekly_report.html"
    template = get_template(template_name)
    context = {
        "reports": response,
        "project": Project.objects.get(id=data.get("project_id")),
        "start_date": end_data.strftime("%d %B, %Y"),
        "end_date": start_date,
    }
    html_content = template.render(context)
    # if doc_type == "docx".lower():
    #     response = HttpResponse(
    #         content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    #     )
    #     response["Content-Disposition"] = 'attachment; filename="weekly_update.docx"'
    #     generate_weekly_update_word(response, context)
    #     return response

    # else:

    #     # Generate PDF
    html = HTML(string=html_content)
    pdf_file = html.write_pdf()
    return HttpResponse(pdf_file, content_type="application/pdf")


# class ProjectListView(ListAPIView):
#     """List all projects with basic information"""
#     queryset = Project.objects.all()
#     serializer_class = ProjectListSerializer
    
#     @swagger_auto_schema(
#         tags=["Case Study"],
#         operation_description="Retrieve a list of all projects",
#         responses={
#             200: ProjectListSerializer(many=True),
#             401: "Unauthorized"
#         }
#     )
#     def get(self, request, *args, **kwargs):
#         return super().get(request, *args, **kwargs)
    
#     def get_queryset(self):
#         return Project.objects.select_related('client').prefetch_related(
#             'platforms', 'categories_tags', 'industries', 'services', 'technology'
#         )


class ProjectPagination(PageNumberPagination):
    page_size = 10
    page_size_query_param = 'page_size'
    max_page_size = 100

class ProjectFilter(django_filters.FilterSet):
    has_services = django_filters.BooleanFilter(
        method='filter_has_services', label='Has Services'
    )
    has_industry = django_filters.BooleanFilter(
        method='filter_has_industry', label='Has Industry'
    )
    has_technology = django_filters.BooleanFilter(
        method='filter_has_technology', label='Has Technology'
    )
    service_ids = django_filters.CharFilter(
        method='filter_by_service_ids', label='Filter by specific service IDs (comma separated)'
    )
    industry_ids = django_filters.CharFilter(
        method='filter_by_industry_ids', label='Filter by specific industry IDs (comma separated)'
    )
    technology_ids = django_filters.CharFilter(
        method='filter_by_technology_ids', label='Filter by specific technology IDs (comma separated)'
    )

    def filter_has_services(self, queryset, name, value):
        if value:
            return queryset.filter(services__isnull=False).distinct()
        return queryset

    def filter_has_industry(self, queryset, name, value):
        if value:
            return queryset.filter(industries__isnull=False).distinct()
        return queryset

    def filter_has_technology(self, queryset, name, value):
        if value:
            return queryset.filter(technology__isnull=False).distinct()
        return queryset

    def filter_by_service_ids(self, queryset, name, value):
        if value:
            try:
                ids = [int(i.strip()) for i in value.split(',') if i.strip()]
                return queryset.filter(services__id__in=ids).distinct()
            except ValueError:
                return queryset.none()
        return queryset

    def filter_by_industry_ids(self, queryset, name, value):
        if value:
            try:
                ids = [int(i.strip()) for i in value.split(',') if i.strip()]
                return queryset.filter(industries__id__in=ids).distinct()
            except ValueError:
                return queryset.none()
        return queryset

    def filter_by_technology_ids(self, queryset, name, value):
        if value:
            try:
                ids = [int(i.strip()) for i in value.split(',') if i.strip()]
                return queryset.filter(technology__id__in=ids).distinct()
            except ValueError:
                return queryset.none()
        return queryset

    class Meta:
        model = Project
        fields = ['has_services', 'has_industry', 'has_technology',
                  'service_ids', 'industry_ids', 'technology_ids']

class ProjectListView(ListAPIView):
    """List all projects with filtering options"""
    queryset = Project.objects.all().order_by('-created_at')
    serializer_class = ProjectListSerializer
    filterset_class = ProjectFilter
    filter_backends = [DjangoFilterBackend, filters.SearchFilter]
    search_fields = ['title', 'client__name', 'description']
    pagination_class = ProjectPagination

    @swagger_auto_schema(
        tags=["Case Study"],
        manual_parameters=[
            openapi.Parameter(
                'has_services', openapi.IN_QUERY,
                description="Filter projects that have at least one service (true/false)",
                type=openapi.TYPE_BOOLEAN
            ),
            openapi.Parameter(
                'has_industry', openapi.IN_QUERY,
                description="Filter projects that have at least one industry (true/false)",
                type=openapi.TYPE_BOOLEAN
            ),
            openapi.Parameter(
                'has_technology', openapi.IN_QUERY,
                description="Filter projects that have at least one technology (true/false)",
                type=openapi.TYPE_BOOLEAN
            ),
            openapi.Parameter(
                'search', openapi.IN_QUERY,
                description="Search across title, client name, and description",
                type=openapi.TYPE_STRING
            ),
            openapi.Parameter(
                'page', openapi.IN_QUERY,
                description="Page number for pagination",
                type=openapi.TYPE_INTEGER
            ),
            openapi.Parameter(
                'page_size', openapi.IN_QUERY,
                description="Number of items per page",
                type=openapi.TYPE_INTEGER
            ),
            openapi.Parameter(
                'service_ids', openapi.IN_QUERY,
                description="Filter projects by specific service IDs (comma separated)",
                type=openapi.TYPE_STRING
            ),
            openapi.Parameter(
                'industry_ids', openapi.IN_QUERY,
                description="Filter projects by specific industry IDs (comma separated)",
                type=openapi.TYPE_STRING
            ),
            openapi.Parameter(
                'technology_ids', openapi.IN_QUERY,
                description="Filter projects by specific technology IDs (comma separated)",
                type=openapi.TYPE_STRING
            ),
        ],
        responses={
            200: openapi.Schema(
                type=openapi.TYPE_OBJECT,
                properties={
                    'count': openapi.Schema(type=openapi.TYPE_INTEGER, description='Total number of projects'),
                    'next': openapi.Schema(type=openapi.TYPE_STRING, description='URL for next page'),
                    'previous': openapi.Schema(type=openapi.TYPE_STRING, description='URL for previous page'),
                    'results': openapi.Schema(
                        type=openapi.TYPE_ARRAY,
                        items=openapi.Schema(
                            type=openapi.TYPE_OBJECT,
                            properties={
                                'id': openapi.Schema(type=openapi.TYPE_INTEGER),
                                'title': openapi.Schema(type=openapi.TYPE_STRING),
                                'slug': openapi.Schema(type=openapi.TYPE_STRING),
                                'description': openapi.Schema(type=openapi.TYPE_STRING),
                                'client_name': openapi.Schema(type=openapi.TYPE_STRING, nullable=True),
                                'client_image': openapi.Schema(type=openapi.TYPE_STRING, nullable=True),
                                'featured_image': openapi.Schema(type=openapi.TYPE_STRING, nullable=True),
                                'thumbnail': openapi.Schema(type=openapi.TYPE_STRING, nullable=True),
                                'live_link': openapi.Schema(type=openapi.TYPE_STRING, nullable=True),
                                'active': openapi.Schema(type=openapi.TYPE_BOOLEAN),
                                'show_in_website': openapi.Schema(type=openapi.TYPE_BOOLEAN),
                                'is_special': openapi.Schema(type=openapi.TYPE_BOOLEAN),
                                'created_at': openapi.Schema(type=openapi.TYPE_STRING, format='date-time'),
                                'updated_at': openapi.Schema(type=openapi.TYPE_STRING, format='date-time'),
                            }
                        ),
                        description='List of filtered projects for current page'
                    ),
                }
            ),
            400: 'Bad Request'
        }
    )
    def get(self, request, *args, **kwargs):
        # Get the filtered queryset
        filtered_queryset = self.filter_queryset(self.get_queryset()).filter(show_in_website=True)
        filtered_count = filtered_queryset.count()

        # Initialize the final queryset
        final_queryset = filtered_queryset

        # If exactly one project is returned, add one is_special=True project
        if filtered_count == 1:
            special_projects = self.get_queryset().filter(
                is_special=True
            ).exclude(
                id__in=filtered_queryset.values('id')
            )[:1]
            # Combine querysets at the Python level
            final_queryset = list(filtered_queryset) + list(special_projects)
        # If no projects are returned, add two is_special=True projects
        elif filtered_count == 0:
            special_projects = self.get_queryset().filter(
                is_special=True
            )[:2]
            final_queryset = list(special_projects)

        # Apply pagination to the final queryset
        page = self.paginate_queryset(final_queryset)
        serializer = self.get_serializer(page, many=True)
        return self.get_paginated_response(serializer.data)

    def get_queryset(self):
        return Project.objects.select_related('client').prefetch_related(
            'platforms', 'categories_tags', 'industries', 'services', 'technology'
        ).filter(show_in_website=True)
    # def get(self, request, *args, **kwargs):
    #     queryset = self.filter_queryset(self.get_queryset()).filter(show_in_website=True)
    #     page = self.paginate_queryset(queryset)
    #     serializer = self.get_serializer(page, many=True)
    #     # print(serializer.data)
    #     return self.get_paginated_response(serializer.data)

    # def get_queryset(self):
    #     return Project.objects.select_related('client').prefetch_related(
    #         'platforms', 'categories_tags', 'industries', 'services', 'technology'
    #     ).filter(show_in_website=True)
    

class ProjectDetailView(RetrieveAPIView):
    """Retrieve project details by slug with all nested data"""
    queryset = Project.objects.all()
    serializer_class = ProjectDetailSerializer
    lookup_field = 'slug'
    
    @swagger_auto_schema(
        tags=["Case Study"],
        operation_description="Retrieve detailed project information by slug",
        manual_parameters=[
            openapi.Parameter(
                'slug',
                openapi.IN_PATH,
                description="Unique slug identifier of the project",
                type=openapi.TYPE_STRING,
                required=True
            )
        ],
        responses={
            200: ProjectDetailSerializer,
            404: "Project not found",
            401: "Unauthorized"
        }
    )
    def get(self, request, *args, **kwargs):
        return super().get(request, *args, **kwargs)
    
    def get_queryset(self):
        # Use the correct relation names based on the error message
        return Project.objects.select_related(
            'client', 'country'
        ).prefetch_related(
            'projectkeypoint_set',
            'projectcontent_set',
            'project_metadata',  # ✅ matches related_name
            'platforms',
            'categories_tags',
            'industries',
            'services',
            'technology',
        )